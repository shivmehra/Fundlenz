from docx import Document

from app.parsers.docx import parse_docx


def test_renders_headings_as_markdown(tmp_path):
    doc = Document()
    doc.add_heading("Fund Overview", level=1)
    doc.add_paragraph("This is the body.")
    doc.add_heading("Performance", level=2)
    doc.add_paragraph("Numbers go here.")
    path = tmp_path / "test.docx"
    doc.save(path)

    out = parse_docx(path)
    assert "# Fund Overview" in out
    assert "## Performance" in out
    assert "This is the body." in out
    assert "Numbers go here." in out


def test_renders_tables_as_markdown(tmp_path):
    doc = Document()
    doc.add_paragraph("Before the table.")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Fund"
    table.rows[0].cells[1].text = "NAV"
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "100"
    table.rows[2].cells[0].text = "B"
    table.rows[2].cells[1].text = "200"
    doc.add_paragraph("After the table.")
    path = tmp_path / "table.docx"
    doc.save(path)

    out = parse_docx(path)
    assert "Before the table." in out
    assert "| Fund | NAV |" in out
    assert "| A | 100 |" in out
    assert "| B | 200 |" in out
    assert "After the table." in out


def test_skips_empty_paragraphs(tmp_path):
    doc = Document()
    doc.add_paragraph("Real content")
    doc.add_paragraph("")
    doc.add_paragraph("More content")
    path = tmp_path / "empty_paras.docx"
    doc.save(path)

    out = parse_docx(path)
    assert "Real content" in out
    assert "More content" in out
    # No runaway blank lines.
    assert "\n\n\n" not in out


def test_preserves_document_order_of_paragraphs_and_tables(tmp_path):
    doc = Document()
    doc.add_heading("Section A", level=1)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "A-cell"
    doc.add_heading("Section B", level=1)
    path = tmp_path / "ordered.docx"
    doc.save(path)

    out = parse_docx(path)
    a_idx = out.index("Section A")
    cell_idx = out.index("A-cell")
    b_idx = out.index("Section B")
    assert a_idx < cell_idx < b_idx


def test_normalizes_unicode_in_paragraphs(tmp_path):
    doc = Document()
    doc.add_paragraph("“Quoted” — dash")  # smart quotes + em dash
    path = tmp_path / "unicode.docx"
    doc.save(path)

    out = parse_docx(path)
    assert '"Quoted" - dash' in out
