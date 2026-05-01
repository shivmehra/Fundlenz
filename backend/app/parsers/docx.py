from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from app.text import normalize


def parse_docx(path: Path) -> str:
    """Extract document text preserving heading hierarchy (rendered as markdown
    `##` headings) and tables (rendered as markdown pipe-tables). Iterates the
    body in document order so paragraphs and tables interleave correctly."""
    doc = Document(path)
    parts: list[str] = []

    for block in _iter_body(doc):
        if block["type"] == "paragraph":
            text = block["text"].strip()
            if not text:
                continue
            level = block["heading_level"]
            if level:
                parts.append(f"{'#' * min(level, 6)} {text}")
            else:
                parts.append(text)
        else:  # table
            md = _docx_table_to_markdown(block["rows"])
            if md:
                parts.append(md)

    return normalize("\n\n".join(parts))


def _iter_body(doc):
    """Yield top-level paragraph and table blocks in document order.

    python-docx's `Paragraph` and `Table` wrappers carry a back-reference to a
    parent that has a `.part` attribute (used for style lookup). Constructing
    them ad hoc from raw body children fails because the body element itself
    has no `.part`. Instead, look up the existing wrapped objects (which were
    constructed with the right parent) by element identity."""
    para_by_id = {id(p._element): p for p in doc.paragraphs}
    table_by_id = {id(t._element): t for t in doc.tables}

    for child in doc.element.body.iterchildren():
        cid = id(child)
        if cid in para_by_id:
            p = para_by_id[cid]
            yield {
                "type": "paragraph",
                "text": p.text,
                "heading_level": _heading_level(p),
            }
        elif cid in table_by_id:
            t = table_by_id[cid]
            rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
            yield {"type": "table", "rows": rows}


def _heading_level(paragraph: Paragraph) -> int:
    """1–9 for `Heading 1`–`Heading 9` styles, 0 for body text."""
    style = paragraph.style.name if paragraph.style else ""
    if style.startswith("Heading "):
        try:
            return int(style.split(" ", 1)[1])
        except (ValueError, IndexError):
            return 0
    return 0


def _docx_table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []
    md = "| " + " | ".join(header) + " |\n"
    md += "| " + " | ".join("---" for _ in header) + " |\n"
    for row in body:
        md += "| " + " | ".join(row) + " |\n"
    return md
